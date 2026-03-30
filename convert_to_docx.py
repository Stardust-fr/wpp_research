#!/usr/bin/env python3
"""
Convert research.txt to formatted DOCX with proper handling of:
- Headings and sections
- Markdown tables to Word tables
- LaTeX equations to formatted text
- Bold, italic, and code formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def parse_markdown_formatting(text):
    """Parse markdown formatting and return runs with proper formatting"""
    # This function returns a list of tuples: (text, bold, italic, code)
    runs = []
    
    # Pattern for bold, italic, code
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            runs.append((part[2:-2], True, False, False))
        elif part.startswith('*') and part.endswith('*'):
            runs.append((part[1:-1], False, True, False))
        elif part.startswith('`') and part.endswith('`'):
            runs.append((part[1:-1], False, False, True))
        else:
            runs.append((part, False, False, False))
    
    return runs

def add_formatted_paragraph(doc, text, style='Normal'):
    """Add paragraph with markdown formatting"""
    para = doc.add_paragraph(style=style)
    
    if not text.strip():
        return para
    
    # Handle LaTeX equations first
    text = text.replace('$$', '\n$$DOUBLEDOLLAR$$\n')
    text = text.replace('$', '\n$SINGLEDOLLAR$\n')
    
    runs_data = parse_markdown_formatting(text)
    
    for text_part, bold, italic, code in runs_data:
        if not text_part:
            continue
        run = para.add_run(text_part)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    return para

def parse_markdown_table(lines, start_idx):
    """Parse markdown table and return list of rows"""
    rows = []
    i = start_idx
    
    while i < len(lines) and lines[i].strip():
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        
        # Split by | and clean up
        cells = [cell.strip() for cell in line.split('|')]
        cells = [cell for cell in cells if cell]  # Remove empty cells
        rows.append(cells)
        i += 1
    
    return rows, i

def add_table_to_doc(doc, rows):
    """Add markdown table to document"""
    if not rows or len(rows) < 2:
        return
    
    # Create table with header row
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Add header row
    header_cells = table.rows[0].cells
    for i, cell_text in enumerate(rows[0]):
        header_cells[i].text = cell_text
        # Format header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
        set_cell_background(header_cells[i], 'D3D3D3')
    
    # Add data rows
    for row in rows[1:]:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            if i < len(row_cells):
                row_cells[i].text = cell_text

def process_research_file(input_file, output_file):
    """Main processing function"""
    doc = Document()
    
    # Read input file
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    current_section_level = 0
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Handle headings
        if line.startswith('## '):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=1)
            i += 1
            continue
        
        if line.startswith('### '):
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=2)
            i += 1
            continue
        
        if line.startswith('#### '):
            heading_text = line[5:].strip()
            doc.add_heading(heading_text, level=3)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip().startswith('---'):
            doc.add_paragraph()
            i += 1
            continue
        
        # Handle markdown tables
        if line.strip().startswith('|'):
            rows, new_i = parse_markdown_table(lines, i)
            if rows:
                add_table_to_doc(doc, rows)
            i = new_i
            continue
        
        # Handle code blocks
        if line.strip().startswith('```'):
            code_block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_block.append(lines[i].rstrip())
                i += 1
            
            if code_block:
                code_para = doc.add_paragraph()
                code_para.paragraph_format.left_indent = Inches(0.5)
                code_para.paragraph_format.right_indent = Inches(0.5)
                code_para.style = 'Quote'
                
                for code_line in code_block:
                    run = code_para.add_run(code_line + '\n')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            
            i += 1  # Skip closing ```
            continue
        
        # Handle bullet points and numbered lists
        if line.strip().startswith('-'):
            para = doc.add_paragraph(line.strip()[2:].strip(), style='List Bullet')
            i += 1
            continue
        
        if line and line[0].isdigit() and '.' in line[:3]:
            match = re.match(r'^\d+\.\s*(.*)', line.strip())
            if match:
                para = doc.add_paragraph(match.group(1), style='List Number')
                i += 1
                continue
        
        # Handle regular paragraphs with formatting
        if line.strip():
            # Handle bold sections like **Metric** at paragraph start
            para_text = line.strip()
            
            # Special handling for lines like "| Metric | Value |" or parameter definitions
            if para_text.startswith('**') and ':' in para_text:
                # This is a key-value pair with bold key
                add_formatted_paragraph(doc, para_text, style='Normal')
            else:
                add_formatted_paragraph(doc, para_text, style='Normal')
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"Successfully created: {output_file}")

if __name__ == '__main__':
    input_file = 'd:\\wpp_research\\research.txt'
    output_file = 'd:\\wpp_research\\WAG-AI_Research_Paper.docx'
    
    process_research_file(input_file, output_file)
